from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DELIVERABLES = ROOT / "entregaveis"
ACTIVITY_SOURCES = ROOT / "diagramas" / "fontes" / "atividades"
DIAGRAM_IMAGES = ROOT / "diagramas" / "imagens"
USE_CASE_SOURCE = ROOT / "diagramas" / "fontes" / "casos-de-uso-stacklyst.puml"

EXPECTED = {
    "AN": {f"AN{i:02d}" for i in range(1, 10)},
    "RF": {f"RF{i:03d}" for i in range(1, 50)},
    "RNF": {f"RNF{i:03d}" for i in range(1, 26)},
    "RN": {f"RN{i:03d}" for i in range(1, 29)},
    "UC": {f"UC{i:03d}" for i in range(1, 26)},
}

PATTERNS = {
    "AN": re.compile(r"\bAN0[1-9]\b"),
    "RF": re.compile(r"\bRF\d{3}\b"),
    "RNF": re.compile(r"\bRNF\d{3}\b"),
    "RN": re.compile(r"(?<!F)\bRN\d{3}\b"),
    "UC": re.compile(r"\bUC\d{3}\b"),
}

LEGACY_PLACEHOLDERS = {
    "[Nome do Projeto]",
    "Descreva aqui",
    "Nome do Requisito",
    "Nome do Caso",
    "Pedro TescaroPedro Tescaro",
    "�",
}


def extract_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join(paragraphs + cells)


def require(condition: bool, message: str, errors: list[str]) -> None:
    if not condition:
        errors.append(message)


def main() -> int:
    errors: list[str] = []
    documents = sorted(DELIVERABLES.glob("*.docx"))
    require(len(documents) == 4, "Devem existir exatamente quatro entregáveis DOCX.", errors)

    combined_text = ""
    for document in documents:
        text = extract_text(document)
        combined_text += "\n" + text
        for placeholder in LEGACY_PLACEHOLDERS:
            require(
                placeholder not in text,
                f"Placeholder antigo encontrado em {document.name}: {placeholder}",
                errors,
            )

    for kind, expected in EXPECTED.items():
        found = set(PATTERNS[kind].findall(combined_text))
        require(
            expected <= found,
            f"Identificadores {kind} ausentes: {sorted(expected - found)}",
            errors,
        )

    activity_sources = sorted(ACTIVITY_SOURCES.glob("AN*.mmd"))
    activity_images = sorted(DIAGRAM_IMAGES.glob("AN*.png"))
    require(len(activity_sources) == 9, "Devem existir nove fontes Mermaid.", errors)
    require(len(activity_images) == 9, "Devem existir nove imagens de atividades.", errors)

    for index in range(1, 10):
        prefix = f"AN{index:02d}"
        require(
            any(path.name.startswith(prefix) for path in activity_sources),
            f"Fonte Mermaid ausente para {prefix}.",
            errors,
        )
        require(
            any(path.name.startswith(prefix) for path in activity_images),
            f"Imagem ausente para {prefix}.",
            errors,
        )

    require(USE_CASE_SOURCE.exists(), "Fonte PlantUML ausente.", errors)
    use_case_image = DIAGRAM_IMAGES / "casos-de-uso-stacklyst.png"
    require(use_case_image.exists(), "Imagem do diagrama de casos de uso ausente.", errors)

    if USE_CASE_SOURCE.exists():
        plantuml = USE_CASE_SOURCE.read_text(encoding="utf-8")
        found_use_cases = set(PATTERNS["UC"].findall(plantuml))
        require(
            EXPECTED["UC"] <= found_use_cases,
            f"Casos de uso ausentes no PlantUML: {sorted(EXPECTED['UC'] - found_use_cases)}",
            errors,
        )

    if errors:
        print("Falha na verificação:")
        for error in errors:
            print(f"- {error}")
        return 1

    print("Verificação concluída com sucesso.")
    print("4 DOCX; 9 atividades; 49 RF; 25 RNF; 28 RN; 25 UC; 10 imagens de diagramas.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
